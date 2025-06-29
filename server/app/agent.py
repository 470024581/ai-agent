import os
from typing import Optional, List, Dict, Any, Union
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits import create_sql_agent
from langchain.agents import AgentType
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from .db import (
    fetch_sales_data_for_query,
    fetch_low_stock_products,
    get_product_details,
    initialize_database,  # Changed: initialize_app_database -> initialize_database
    get_files_by_datasource, # Added to get files for RAG
    DATABASE_PATH # Import DATABASE_PATH
)
from .report import generate_daily_sales_summary_report
from .models import DataSourceType # Import DataSourceType
from dotenv import load_dotenv
from pathlib import Path # Added Path
import logging # Added logging
import pandas as pd # Added pandas
import re
import asyncio
from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import for local embeddings
# from langchain_community.embeddings import SentenceTransformerEmbeddings # Already imported above

load_dotenv()

# Factory imports
from .llm_factory import get_llm, get_llm_status, reset_llm
from .embedding_factory import get_embeddings, get_embeddings_status, reset_embeddings

# Determine the correct upload directory relative to this file (agent.py)
# Assuming agent.py is in server/app/ and uploads are in server/data/uploads/
UPLOAD_DIR = Path(__file__).resolve().parent.parent / "data" / "uploads"
DB_URI = f"sqlite:///{DATABASE_PATH}" # Construct DB URI for LangChain

llm = None
embeddings = None # Initialize embeddings variable

# Imports for file parsing
import PyPDF2
from docx import Document as DocxDocument

# Helper functions for parsing files
def _extract_text_from_pdf(file_path: Path) -> str:
    logger.info(f"Extracting text from PDF: {file_path}")
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() or ""
        logger.info(f"Successfully extracted {len(text)} characters from PDF: {file_path}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
    return text

def _extract_text_from_docx(file_path: Path) -> str:
    logger.info(f"Extracting text from DOCX: {file_path}")
    text = ""
    try:
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        logger.info(f"Successfully extracted {len(text)} characters from DOCX: {file_path}")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
    return text

def _extract_text_from_csv_pandas(file_path: Path) -> str:
    logger.info(f"Extracting text from CSV using pandas: {file_path}")
    text = ""
    try:
        df = pd.read_csv(file_path, on_bad_lines='skip')
        # Convert each row to a string, then join all row strings
        # We include column names for context for each row, and join with a clear separator.
        # For example: "column1: value1, column2: value2, ..."
        row_texts = []
        for index, row in df.iterrows():
            row_text = ", ".join([f"{col}: {str(val)}" for col, val in row.astype(str).items()])
            row_texts.append(row_text)
        text = "\n".join(row_texts) # Each original row becomes a line in the text document
        logger.info(f"Successfully extracted text from CSV {file_path}. Total characters: {len(text)}")
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path} with pandas: {e}", exc_info=True)
    return text

def _extract_text_from_xlsx_pandas(file_path: Path) -> str:
    logger.info(f"Extracting text from XLSX using pandas: {file_path}")
    text = ""
    try:
        xls = pd.ExcelFile(file_path)
        sheet_texts = []
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            # Convert each row to a string, then join all row strings for this sheet
            row_texts = []
            for index, row in df.iterrows():
                row_text = ", ".join([f"{col}: {str(val)}" for col, val in row.astype(str).items()])
                row_texts.append(row_text)
            sheet_texts.append(f"Sheet: {sheet_name}\n" + "\n".join(row_texts))
        text = "\n\n".join(sheet_texts) # Separate sheets by a double newline
        logger.info(f"Successfully extracted text from XLSX {file_path}. Total characters: {len(text)}")
    except Exception as e:
        logger.error(f"Error extracting text from XLSX {file_path} with pandas: {e}", exc_info=True)
    return text

# Initialize LLM using factory (strict mode - no simulation)
try:
    llm = get_llm()
    llm_status = get_llm_status()
    logger.info(f"LLM initialized successfully - Provider: {llm_status.get('provider')}, Model: {llm_status.get('model')}")
    if llm_status.get('base_url'):
        logger.info(f"Using API endpoint: {llm_status.get('base_url')}")
except Exception as e:
    logger.error(f"LLM initialization failed: {e}")
    logger.error("Please check your LLM configuration in .env file")
    llm = None

# Initialize Embeddings using factory
try:
    logger.info("Starting initialization of embedding model using factory...")
    embeddings = get_embeddings()
    embedding_status = get_embeddings_status()
    logger.info(f"Embedding model initialized successfully - Provider: {embedding_status.get('provider')}, Model: {embedding_status.get('model')}")
except Exception as e:
    logger.error(f"Embedding model initialization failed: {e}", exc_info=True)
    embeddings = None

async def perform_rag_query(query: str, datasource: Dict[str, Any]) -> Dict[str, Any]:
    """
    Performs RAG retrieval and Q&A for the specified data source.
    """
    logger.info(f"Attempting RAG query on datasource: {datasource['name']} (ID: {datasource['id']}) for query: '{query}'")

    if not llm:
        raise RuntimeError("LLM not initialized. RAG query cannot be processed without LLM.")
    
    if not embeddings:
        logger.error("Local embedding model not initialized. RAG query cannot perform vectorization and retrieval.")
        return {
            "query": query, "query_type": "rag", "success": False,
            "answer": "RAG functionality cannot be executed because a dependent component is not initialized. Please check API keys and configuration.",
            "data": {"source_datasource_id": datasource['id'], "source_datasource_name": datasource['name']},
            "error": "Local embeddings not initialized."
        }
    
    try:
        # 1. Fetch list of 'completed' files from the database
        datasource_id = datasource['id']
        logger.info(f"Fetching completed files for datasource_id: {datasource_id}")
        db_files = await get_files_by_datasource(datasource_id)
        
        completed_files = [f for f in db_files if f['processing_status'] == 'completed']
        logger.info(f"Found {len(completed_files)} completed files for datasource {datasource_id}.")

        if not completed_files:
            return {
                "query": query, "query_type": "rag", "success": True, # Success=True as it's a valid state
                "answer": f"No successfully processed files available for query in data source '{datasource['name']}'. Please upload files and wait for processing to complete.",
                "data": {"source_datasource_id": datasource['id'], "source_datasource_name": datasource['name'], "retrieved_documents": []}
            }

        all_docs = []
        # 2. Load and parse file content
        for file_info in completed_files:
            file_path = UPLOAD_DIR / file_info['filename']
            text_content = ""
            file_type = file_info['file_type']
            original_filename = file_info['original_filename']

            logger.info(f"Processing file: {file_path} (Original: {original_filename}, Type: {file_type})")

            if not file_path.exists():
                logger.warning(f"File not found: {file_path} for file_info: {original_filename}")
                continue
            
            try:
                if file_type == 'txt':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                elif file_type == 'pdf':
                    text_content = _extract_text_from_pdf(file_path)
                elif file_type == 'docx':
                    text_content = _extract_text_from_docx(file_path)
                # Add other file types like CSV, XLSX later
                elif file_type == 'csv':
                    text_content = _extract_text_from_csv_pandas(file_path)
                elif file_type == 'xlsx':
                    text_content = _extract_text_from_xlsx_pandas(file_path)
                else:
                    logger.info(f"Skipping file {original_filename} due to unsupported file type: {file_type}")
                    continue
                
                if text_content.strip():
                    doc = Document(page_content=text_content, metadata={"source": original_filename, "file_id": file_info['id']})
                    all_docs.append(doc)
                    logger.info(f"Successfully processed and created Document for {original_filename}. Length: {len(text_content)}")
                else:
                    logger.warning(f"No text content extracted from {original_filename} (Type: {file_type})")

            except Exception as e:
                logger.error(f"Error processing file {original_filename}: {e}", exc_info=True)
                # Optionally, update file status to 'failed' here if processing fails at this stage
                # await update_file_processing_status(file_info['id'], status=ProcessingStatus.FAILED.value, error_message=f"Content extraction failed: {str(e)}")

        if not all_docs:
             return {
                "query": query, "query_type": "rag", "success": True,
                "answer": f"No processable (TXT, PDF, DOCX, CSV, XLSX) file content found in data source '{datasource['name']}'.",
                "data": {"source_datasource_id": datasource['id'], "source_datasource_name": datasource['name'], "retrieved_documents": []}
            }

        # 3. Text chunking
        logger.info(f"Splitting {len(all_docs)} documents into chunks.")
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunked_docs = text_splitter.split_documents(all_docs)
        logger.info(f"Created {len(chunked_docs)} chunks.")

        if not chunked_docs:
            logger.warning("No chunks were created from the documents. Cannot proceed with RAG.")
            return {
                "query": query, "query_type": "rag", "success": True,
                "answer": f"Could not extract valid content chunks from files in data source '{datasource['name']}' for querying.",
                "data": {"source_datasource_id": datasource['id'], "source_datasource_name": datasource['name'], "retrieved_documents": []}
            }

        # 4. Create in-memory vector store (FAISS)
        logger.info("Creating FAISS vector store from chunks...")
        vector_store = FAISS.from_documents(chunked_docs, embeddings)
        logger.info("FAISS vector store created.")

        # 5. Perform retrieval (RetrievalQA chain)
        logger.info("Setting up RetrievalQA chain...")
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff", # Other types: map_reduce, refine, map_rerank
            retriever=vector_store.as_retriever(search_kwargs={"k": 3}), # Retrieve top 3 chunks
            return_source_documents=True
        )
        logger.info(f"Executing RAG query: '{query}'")
        result = await qa_chain.ainvoke({"query": query})
        
        answer = result.get("result", "Could not find a clear answer in the knowledge base.")
        source_documents_data = []
        if result.get("source_documents"):
            for doc_source in result["source_documents"]:
                source_documents_data.append({
                    "source": doc_source.metadata.get("source", "Unknown source"),
                    "content_preview": doc_source.page_content[:200] + "..." # Preview of content
                })

        logger.info(f"RAG query successful. Answer: {answer[:100]}... Sources: {len(source_documents_data)}")
        return {
            "query": query,
            "query_type": "rag",
            "answer": answer,
            "data": {
                "source_datasource_id": datasource['id'],
                "source_datasource_name": datasource['name'],
                "retrieved_documents": source_documents_data
            },
            "success": True
        }

    except Exception as e:
        logger.error(f"Error during RAG for datasource {datasource['name']}: {e}", exc_info=True)
        return {
            "query": query, "query_type": "rag", "success": False,
            "answer": f"A critical error occurred while processing your query '{query}' in data source '{datasource['name']}'.",
            "data": {"source_datasource_id": datasource['id'], "source_datasource_name": datasource['name']},
            "error": str(e)
        }

async def get_answer_from_hybrid_datasource(query: str, active_datasource: Dict[str, Any]) -> Dict[str, Any]:
    """
    Handles queries for hybrid data sources.
    Analyzes the query and available files to determine whether to use SQL or RAG processing.
    
    Decision logic:
    1. If the datasource has SQL tables (db_table_name) and query seems structured -> SQL
    2. If query seems document-related or no SQL tables available -> RAG
    3. Can potentially combine both approaches for complex queries
    """
    logger.info(f"Processing hybrid query for datasource: {active_datasource['name']} (ID: {active_datasource['id']})")
    
    datasource_id = active_datasource['id']
    ds_name = active_datasource['name']
    
    try:
        # Get available files and their types
        db_files = await get_files_by_datasource(datasource_id)
        completed_files = [f for f in db_files if f['processing_status'] == 'completed']
        
        # Categorize files by processing type
        sql_files = [f for f in completed_files if f['file_type'].lower() in ['csv', 'xlsx']]
        rag_files = [f for f in completed_files if f['file_type'].lower() in ['txt', 'pdf', 'docx']]
        
        logger.info(f"Hybrid datasource analysis - SQL files: {len(sql_files)}, RAG files: {len(rag_files)}")
        
        # Determine query routing strategy
        query_lower = query.lower()
        
        # Keywords that suggest SQL operations
        sql_keywords = ['sum', 'count', 'average', 'total', 'sales', 'revenue', 'amount', 'quantity', 
                       'statistics', 'report', 'calculate', 'how many', 'how much', 'trend', 'analysis']
        
        # Keywords that suggest document/RAG operations  
        rag_keywords = ['what is', 'explain', 'describe', 'definition', 'meaning', 'content', 
                       'document', 'text', 'information about', 'tell me about', 'summary of']
        
        has_sql_indicators = any(keyword in query_lower for keyword in sql_keywords)
        has_rag_indicators = any(keyword in query_lower for keyword in rag_keywords)
        
        # Decision logic
        if sql_files and active_datasource.get("db_table_name") and (has_sql_indicators or not has_rag_indicators):
            # Route to SQL if we have SQL data and query seems numerical/analytical
            logger.info(f"Routing hybrid query to SQL processing for datasource: {ds_name}")
            return await get_answer_from_sqltable_datasource(query, active_datasource)
            
        elif rag_files and (has_rag_indicators or not has_sql_indicators):
            # Route to RAG if we have documents and query seems informational
            logger.info(f"Routing hybrid query to RAG processing for datasource: {ds_name}")
            return await perform_rag_query(query, active_datasource)
            
        elif sql_files and active_datasource.get("db_table_name"):
            # Default to SQL if available
            logger.info(f"Defaulting hybrid query to SQL processing for datasource: {ds_name}")
            return await get_answer_from_sqltable_datasource(query, active_datasource)
            
        elif rag_files:
            # Default to RAG if SQL not available
            logger.info(f"Defaulting hybrid query to RAG processing for datasource: {ds_name}")
            return await perform_rag_query(query, active_datasource)
            
        else:
            # No processed files available
            return {
                "query": query, "query_type": "hybrid", "success": True,
                "answer": f"Hybrid data source '{ds_name}' doesn't have any processed files available. Please upload CSV/Excel files for SQL queries or TXT/PDF/Word files for document queries.",
                "data": {
                    "source_datasource_id": datasource_id,
                    "source_datasource_name": ds_name,
                    "available_sql_files": len(sql_files),
                    "available_rag_files": len(rag_files),
                    "routing_decision": "no_files_available"
                }
            }
            
    except Exception as e:
        logger.error(f"Error in hybrid datasource processing: {e}", exc_info=True)
        return {
            "query": query, "query_type": "hybrid", "success": False,
            "answer": f"Error processing hybrid query for data source '{ds_name}': {str(e)}",
            "data": {"source_datasource_name": ds_name},
            "error": str(e)
        }

async def get_answer_from_sqltable_datasource(query: str, active_datasource: Dict[str, Any]) -> Dict[str, Any]:
    """
    Queries the dynamically created SQL table associated with the specified data source using LangChain SQL Agent.
    """
    logger.info(f"Attempting SQL Agent query on datasource: {active_datasource['name']} (ID: {active_datasource['id']}) for query: '{query}'")

    if not llm:
        logger.error("LLM not initialized. SQL Agent query cannot be performed.")
        return {
            "query": query, "query_type": "sql_agent", "success": False,
            "answer": "SQL Agent functionality cannot be executed because LLM is not initialized.",
            "data": {"source_datasource_id": active_datasource['id'], "source_datasource_name": active_datasource['name']},
            "error": "LLM not initialized for SQL Agent."
        }

    db_table_name = active_datasource.get("db_table_name")
    if not db_table_name:
        logger.error(f"Data source '{active_datasource['name']}' is not fully configured. Missing associated database table name.")
        return {
            "query": query, "query_type": "sql_agent", "success": False,
            "answer": f"Data source '{active_datasource['name']}' is not fully configured. Missing associated database table name.",
            "data": {"source_datasource_id": active_datasource['id'], "source_datasource_name": active_datasource['name']},
            "error": "Missing db_table_name for SQL_TABLE_FROM_FILE datasource."
        }

    try:
        logger.info(f"Initializing SQLDatabase for table: {db_table_name} using URI: {DB_URI}")
        # SQLDatabase will connect to the main smart.db, but we tell it to only include the specific table.
        db = SQLDatabase.from_uri(DB_URI, include_tables=[db_table_name], sample_rows_in_table_info=3)
        
        # Get schema info for validation
        schema_info = db.get_table_info()
        
        # Process time expressions in query
        processed_query = _process_time_expressions(query)

        # Get available tables
        available_tables = db.get_usable_table_names()
        logger.info(f"Available tables: {available_tables}")

        # Detect if this is a time series/trend query
        is_trend_query = any(keyword in processed_query.lower() for keyword in [
            'trend', 'monthly', 'weekly', 'yearly', 'over time', 'by month', 'by year', 
            'time series', 'sales trend', '2025', 'quarterly'
        ])

        # Generate SQL using direct LLM call with enhanced time series support
        sql_generation_prompt = f"""
        User query: "{processed_query}"

        Available tables: {available_tables}
        Table schema: {schema_info}

        **CRITICAL ANALYSIS**: This query appears to be a {'TIME SERIES/TREND' if is_trend_query else 'STANDARD'} query.

        Please generate a SQLite query to answer the user's question. The query should:
        1. Only use SELECT statements (no INSERT, UPDATE, DELETE, etc.)
        2. Only use tables that are available in the database
        
        **FOR TIME SERIES/TREND QUERIES** (like "monthly sales trend", "sales by month", etc.):
        3a. Use time grouping with strftime functions:
           - Monthly trends: SELECT strftime('%Y-%m', saledate) as month, SUM(totalamount) as total_sales
           - Weekly trends: SELECT strftime('%Y-W%W', saledate) as week, SUM(totalamount) as total_sales  
           - Yearly trends: SELECT strftime('%Y', saledate) as year, SUM(totalamount) as total_sales
        3b. Include GROUP BY with the same time function
        3c. Order by the time field for proper chronological display
        3d. Example for monthly trend: 
            SELECT strftime('%Y-%m', saledate) as month, SUM(totalamount) as total_sales 
            FROM table_name 
            WHERE strftime('%Y', saledate) = '2025'
            GROUP BY strftime('%Y-%m', saledate) 
            ORDER BY strftime('%Y-%m', saledate)
            
        **FOR STANDARD QUERIES** (like "top products", "product sales", etc.):
        3e. Use appropriate filters and aggregations based on the question
        3f. Use time filters for temporal queries:
           - "this month" -> WHERE strftime('%Y-%m', saledate) = strftime('%Y-%m', 'now')
           - "last month" -> WHERE strftime('%Y-%m', saledate) = strftime('%Y-%m', 'now', '-1 month')
           - "this year" -> WHERE strftime('%Y', saledate) = strftime('%Y', 'now')
           
        4. Use appropriate aggregations (SUM, COUNT, AVG) based on what user asks
        5. Include proper WHERE clauses to filter data as needed
        6. Order results appropriately (time series by time, others by value DESC)  
        7. Limit results to a reasonable number (e.g., LIMIT 20 for trends, LIMIT 10 for products)

        **IMPORTANT**: 
        - For trend queries, return (time_period, aggregated_value) pairs
        - For product queries, return (product_name, aggregated_value) pairs
        - Match the query intent precisely

        Only return the SQL query, no explanations or other text.
        """

        try:
            # Execute query with timeout from config
            sql_response = await asyncio.wait_for(
                llm.agenerate([sql_generation_prompt]),
                timeout=Config.LLM_TIMEOUT
            )
            
            # Extract SQL from response
            sql = sql_response.generations[0][0].text.strip()
            
            # Clean and validate SQL
            clean_sql = _clean_sql_statement(sql)
            if not _validate_sql_statement(clean_sql, schema_info):
                raise ValueError(f"Invalid SQL statement generated: {clean_sql}")
            
            try:
                # Execute the cleaned SQL
                query_result = db.run(clean_sql)
                
                # Handle different result types
                if isinstance(query_result, str):
                    # If result is a string, it might be an error or a single value
                    logger.warning(f"Query returned string result: {query_result}")
                    
                    # Try to parse if it's a string representation of data
                    parsed_data = _parse_string_query_result(query_result)
                    if parsed_data:
                        structured_data = {
                            "rows": parsed_data,
                            "columns": ["label", "value"] if len(parsed_data[0]) == 2 else [f"col_{i}" for i in range(len(parsed_data[0]))],
                            "executed_sql": clean_sql,
                            "queried_table": db_table_name
                        }
                    else:
                        # Fallback to original logic for simple string results
                        structured_data = {
                            "rows": [{"result": query_result}],
                            "columns": ["result"],
                            "executed_sql": clean_sql,
                            "queried_table": db_table_name
                        }
                else:
                    # Normal case: result is a list of dictionaries
                    if not query_result:
                        return {
                            "success": False,
                            "error": "Query returned no results. Please check if the time period contains data.",
                            "query": processed_query,
                            "executed_sql": clean_sql
                        }
                        
                    # Ensure query_result is a list
                    if not isinstance(query_result, list):
                        query_result = [query_result]
                        
                    structured_data = {
                        "rows": query_result,
                        "columns": list(query_result[0].keys()) if query_result else [],
                        "executed_sql": clean_sql,
                        "queried_table": db_table_name
                    }
                
                return {
                    "success": True,
                    "data": structured_data,
                    "answer": f"Here are the results for your query about {processed_query}",
                    "executed_sql": clean_sql,
                    "query": processed_query
                }
                
            except Exception as e:
                error_msg = str(e)
                if "no such table" in error_msg.lower():
                    error_msg = f"Table not found. Available tables are: {', '.join(available_tables)}"
                elif "no such column" in error_msg.lower():
                    error_msg = f"Column not found in schema: {error_msg}"
                    
                logger.error(f"SQL execution failed: {error_msg}")
                return {
                    "success": False,
                    "error": f"Error executing SQL query: {error_msg}",
                    "query": processed_query,
                    "executed_sql": clean_sql
                }

        except asyncio.TimeoutError:
            logger.warning(f"Query SQL generation timed out for query: {processed_query}")
            return {
                "success": False,
                "error": "Query generation timed out",
                "query": processed_query
            }

    except Exception as e:
        logger.error(f"Error during SQL Agent query: {e}")
        return {
            "query": query, "query_type": "sql_agent", "success": False,
            "answer": f"Error during SQL Agent query: {str(e)}",
            "data": {"source_datasource_id": active_datasource['id'], "source_datasource_name": active_datasource['name']},
            "error": str(e)
        }

async def get_answer_from(query: str, query_type: str = "sales", active_datasource: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Main function to get answers from the  system.
    Routes the query based on the query_type and active_datasource.
    """
    logger.info(f"get_answer_from called with query: '{query}', query_type: '{query_type}', active_datasource: {active_datasource['name'] if active_datasource else 'None'}")

    if not llm and (not active_datasource or active_datasource.get('type') != DataSourceType.DEFAULT.value):
         # If LLM is not available AND we are not using the default  (which might have non-LLM paths)
        raise RuntimeError("LLM not initialized. Cannot process query without LLM.")
        # Allow default  queries to proceed if they don't rely on LLM for their basic logic.
        # For RAG or SQL Agent, LLM is crucial.
        if active_datasource and active_datasource.get('type') != DataSourceType.DEFAULT.value:
            return {
                "query": query, "query_type": query_type, "success": False,
                "answer": "Core question answering service not initialized. Unable to process your request. Please check system configuration.",
                "data": {"source_datasource_name": active_datasource['name'] if active_datasource else "N/A"},
                "error": "LLM not initialized."
            }

    # Determine the type of the active datasource
    ds_type = active_datasource.get('type') if active_datasource else DataSourceType.DEFAULT.value
    ds_name = active_datasource.get('name', 'Default ') if active_datasource else 'Default '

    logger.info(f"Routing query. Datasource type: {ds_type}, Datasource name: {ds_name}")

    if ds_type == DataSourceType.SQL_TABLE_FROM_FILE.value:
        logger.info(f"Routing to SQL Agent for datasource: {ds_name}")
        if not active_datasource or not active_datasource.get("db_table_name"):
            logger.error(f"SQL_TABLE_FROM_FILE datasource '{ds_name}' is missing 'db_table_name'.")
            return {
                "query": query, "query_type": "sql_agent", "success": False,
                "answer": f"Data source '{ds_name}' is not fully configured. Unable to perform SQL query.",
                "data": {"source_datasource_name": ds_name},
                "error": "Missing db_table_name."
            }
        return await get_answer_from_sqltable_datasource(query, active_datasource)
    
    elif ds_type == DataSourceType.HYBRID.value:
        logger.info(f"Routing to Hybrid Agent for datasource: {ds_name}")
        return await get_answer_from_hybrid_datasource(query, active_datasource)
    
    elif ds_type == DataSourceType.KNOWLEDGE_BASE.value: # Knowledge base uses RAG
        logger.info(f"Routing to RAG for datasource: {ds_name}")
        if not active_datasource: # Should not happen if ds_type is not DEFAULT, but good check
             return {"answer": "Error: Attempting to use RAG without specifying a custom data source."}
        return await perform_rag_query(query, active_datasource)

    # Fallback to default  logic (SQLite queries for products, sales, inventory)
    logger.info(f"Routing to default  logic for datasource: {ds_name} (Query Type Hint: {query_type})")
    
    # Existing  logic based on query_type (sales, inventory, report)
    # This part might need LLM for interpretation or can use direct DB queries for simple cases.
    # For now, it assumes direct DB queries or simplified LLM interaction.
    if query_type == "sales":
        answer, data = await get_sales_query_response(query)
        return {"answer": answer, "data": data, "query_type": "sales", "source_datasource_name": ds_name}
    elif query_type == "inventory":
        data, answer = await get_inventory_check_response_by_query(query)
        return {"answer": answer, "data": {"low_stock_items": data}, "query_type": "inventory", "source_datasource_name": ds_name}
    elif query_type == "report" or "report" in query or "summary" in query:
        answer, report_data, chart_data = await get_daily_sales_report_response_from_agent()
        return {"answer": answer, "data": report_data, "chart_data": chart_data, "query_type": "report", "source_datasource_name": ds_name}
    else: # Default/Fallback for  if query_type is not specific
        # This could be a generic LLM call against the schema or a predefined set of actions.
        # For simplicity, let's route to sales as a general default for now.
        logger.info(f"Default  query type not specific ('{query_type}'), defaulting to sales-like query.")
        answer, data = await get_sales_query_response(query) # Or a more generic  query handler
        return {"answer": answer, "data": data, "query_type": "general", "source_datasource_name": ds_name}


async def get_sales_query_response(query: str) -> tuple[str, Optional[dict]]:
    """
    Handles sales-related queries for the default  datasource.
    (This is part of the original  logic)
    """
    logger.info(f"Processing sales query for default : {query}")
    # This function would contain logic to parse the query, fetch data from 
    # the 'sales', 'products' tables in SQLite, and formulate an answer.
    # It might use LLM for understanding the query and formatting the response,
    # or use simpler rule-based logic for predefined questions.
    
    # Example: Fetch raw data first
    sales_data = await fetch_sales_data_for_query(query) # This function already exists

    if not sales_data:
        return "Based on your query, I couldn't find related sales data.", None

    # Simplified response for now. A real scenario might use LLM to summarize.
    if llm:
        # Prepare a prompt for the LLM to summarize the sales_data based on the query
        prompt = f"""
        User query: "{query}"
        Relevant sales data (JSON):
        {str(sales_data[:5])} ... (showing first 5 records for brevity)
        
        Please generate a concise answer in English based on the user query and the data above.
        If the user query asks for specific metrics (e.g., total sales, average order value), calculate and include them.
        """
        try:
            response = llm.invoke(prompt)
            answer = response.content
            logger.info(f"LLM generated sales answer: {answer}")
            return answer, {"detailed_sales": sales_data}
        except Exception as e:
            logger.error(f"LLM invocation failed for sales query summarization: {e}", exc_info=True)
            return f"Found {len(sales_data)} related sales records. Unable to provide detailed summary due to LLM processing error.", {"detailed_sales": sales_data}
    else:
        # Fallback if LLM is not available for sales query
        num_records = len(sales_data)
        total_sales_amount = sum(item['total_amount'] for item in sales_data)
        
        answer = f"Found {num_records} sales records. Total sales amount approximately {total_sales_amount:.2f}."
        if "best selling" in query or "top selling" in query:
            # Simple logic for top selling - could be more sophisticated
            from collections import Counter
            product_counts = Counter(item['product_name'] for item in sales_data)
            top_product, count = product_counts.most_common(1)[0] if product_counts else ("Unknown", 0)
            answer += f" The most frequently sold product is '{top_product}' (sold {count} times)."

        return answer, {"detailed_sales": sales_data, "summary_stats": {"total_records": num_records, "total_amount": total_sales_amount}}

async def get_inventory_check_response_by_query(query: str) -> tuple[List[Dict[str, Any]], str]:
    """
    Handles inventory-related queries by trying to extract a product ID or using general low stock.
    (This is part of the original  logic)
    """
    logger.info(f"Processing inventory query for default : {query}")
            # Simple check for a product ID (e.g., "inventory P001") - this could be more robust
    # For a more general query, it might default to showing low stock items.
    
    # Attempt to find a product ID in the query (very basic)
    # A more robust solution would use NER or pattern matching.
    potential_product_id = None
    words = query.split()
    for word in words:
        if word.upper().startswith("P") and word[1:].isdigit(): # Simple check for "PXXX" format
            potential_product_id = word.upper()
            break
            
    items_to_report = []
    response_summary = ""

    if potential_product_id:
        logger.info(f"Found potential product ID in inventory query: {potential_product_id}")
        product_detail = await get_product_details(potential_product_id) # Assumes get_product_details also fetches inventory
        if product_detail:
            # This assumes get_product_details would be extended to include stock info
            # or we make another call here. For now, let's assume it's part of product_detail.
            # items_to_report.append(product_detail) # Adjust based on actual structure
            # response_summary = f"Product {product_detail.get('product_name', potential_product_id)} inventory information..."
            
            # Let's refine this - get_product_details gets product info, then we need inventory for it.
            # This part needs to be aligned with how inventory is fetched per product.
            # For now, let's simulate:
            # Assume we have a function like `get_inventory_for_product(product_id)`
            # This is a placeholder - actual implementation would query inventory table.
            # For demonstration, we'll just say "Details for PXXX..." and rely on general low stock if not specific.
            # This logic is a bit convoluted due to missing specific inventory function per product.
            # Let's simplify: if specific product ID, try to get its details. Otherwise, low stock.
            
            # A better approach for "inventory of P001" is to make `get_product_details` more comprehensive
            # or have a dedicated `get_inventory_for_product(product_id)` function.
            # Given the current `get_low_stock_products`, we might not have a direct single product inventory view easily.
            
            # Let's assume the query "inventory P001" means "is P001 low stock?" for simplicity with current tools.
            low_stock_items = await fetch_low_stock_products(threshold=1000) # High threshold to get most items
            found_product_in_low_stock = False
            for item in low_stock_items:
                if item['product_id'] == potential_product_id:
                    items_to_report.append(item)
                    response_summary = f"Product {item['product_name']} (ID: {item['product_id']}) current stock level is {item['stock_level']}."
                    if item['stock_level'] < 50: # Example threshold
                        response_summary += " Stock level is low."
                    else:
                        response_summary += " Stock is sufficient."
                    found_product_in_low_stock = True
                    break
            if not found_product_in_low_stock:
                # If not in low_stock_items, assume it's not low or doesn't exist.
                # A full check would query inventory table directly.
                # For now, we say:
                response_summary = f"No low stock record found for product ID {potential_product_id}. Stock may be sufficient or product does not exist."
                # To be more accurate, we'd need a direct inventory lookup.
        else:
            response_summary = f"No product information found for product ID {potential_product_id}."

    else:  # General inventory query (corresponds to if potential_product_id)
        logger.info("General inventory query, fetching low stock items.")
        low_stock_items = await fetch_low_stock_products(threshold=50) # Default threshold
        if low_stock_items:
            items_to_report = low_stock_items
            response_summary = f"Currently {len(low_stock_items)} products have low stock (below 50 units). Please replenish."
            # For brevity, we don't list them all in the summary string here, they are in `data`.
        else: # Corresponds to if low_stock_items
            response_summary = "Currently all products have stock above the warning threshold (50 units)."
            # items_to_report will remain empty
            
    # If LLM is available, it could rephrase `response_summary` or interpret `items_to_report`
    if llm:
        prompt = f"""
        User inventory query: "{query}"
        Retrieved inventory information/summary: "{response_summary}"
        Relevant product data (JSON):
        {str(items_to_report[:3])} ... (showing first 3 items for brevity if many)

        Please generate a natural answer in English based on the user query and the data above.
        """
        try:
            llm_response = llm.invoke(prompt)
            final_answer = llm_response.content
            logger.info(f"LLM generated inventory answer: {final_answer}")
            return items_to_report, final_answer
        except Exception as e:
            logger.error(f"LLM invocation failed for inventory query summarization: {e}", exc_info=True)
            # Fallback to pre-LLM summary if LLM fails during try block
            return items_to_report, response_summary
    else: # This is the corrected else for 'if llm:'
        return items_to_report, response_summary

async def get_inventory_check_response() -> tuple[List[Dict[str, Any]], str]:
    # This is a simplified version, assuming a general "check inventory" means "show low stock"
    # Kept for backward compatibility or direct calls if any.
    low_stock_items = await fetch_low_stock_products(threshold=50)
    if low_stock_items:
        answer = f"Found {len(low_stock_items)} low stock products (below 50 units):"
        # for item in low_stock_items:
        #     answer += f"\n- {item['product_name']} (ID: {item['product_id']}), current stock: {item['stock_level']}"
    else:
        answer = "All products have stock above the warning threshold (50 units)."
    return low_stock_items, answer

async def get_daily_sales_report_response_from_agent() -> tuple[str, Dict[str, Any], Optional[Dict[str, Any]]]:
    """
    Generates a daily sales report summary using the report generation utility.
    (This is part of the original  logic)
    """
    logger.info("Generating daily sales report for default .")
    # This uses the existing report generation logic.
    # If LLM is available, it could enhance the summary.
    summary, report_data, chart_data = await generate_daily_sales_summary_report(llm=llm) # Pass LLM if needed by report
    
    # The summary from generate_daily_sales_summary_report might already be LLM-generated or template-based.
    # If not, and LLM is available here, we could try to enhance it.
    # For now, we assume the summary is good as is.
    
    return summary, report_data, chart_data

async def attempt_direct_query_fallback(query: str, db, table_name: str, active_datasource: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    Fallback solution: execute simple database queries directly to avoid Agent parsing issues
    """
    try:
        # Generate simple SQL based on query content
        if "sales" in query.lower() or "trend" in query.lower():
            # Sales trend query
            direct_sql = f"SELECT date, sales FROM {table_name} ORDER BY date LIMIT 10"
        elif "product" in query.lower():
            # Product query
            direct_sql = f"SELECT product, SUM(sales) as total_sales FROM {table_name} GROUP BY product LIMIT 10"
        elif "total" in query.lower():
            # Total query
            direct_sql = f"SELECT SUM(sales) as total_sales FROM {table_name}"
        else:
            # Default query
            direct_sql = f"SELECT * FROM {table_name} LIMIT 10"
        
        logger.info(f"Executing direct fallback query: {direct_sql}")
        result = db.run(direct_sql)
        
        if result:
            # Parse results and generate answer
            if "trend" in query.lower():
                answer = "The sales trend shows data from the database query results."
            elif "total" in query.lower():
                answer = f"The total sales amount is {result}."
            else:
                answer = f"Found {len(result.split(',')) if isinstance(result, str) else 'several'} records based on your query."
            
            return {
                "query": query,
                "query_type": "sql_agent",
                "success": True,
                "answer": answer,
                "data": {
                    "source_datasource_id": active_datasource['id'],
                    "source_datasource_name": active_datasource['name'],
                    "queried_table": table_name,
                    "rows": result,
                    "recovery_method": "direct_query_fallback"
                }
            }
    except Exception as e:
        logger.error(f"Direct query fallback failed: {e}")
        return None

def extract_answer_from_timeout(timeout_response: str, original_query: str, db, table_name: str) -> str:
    """
    Try to extract useful information from timeout Agent response, or execute simple direct query
    """
    try:
        from sqlalchemy import text
        # Try executing a simple table data query as fallback solution
        simple_query = f"SELECT * FROM {table_name} LIMIT 5"
        
        with db._engine.connect() as connection:
            result = connection.execute(text(simple_query))
            rows = result.fetchall()
            columns = result.keys()
            
            if rows:
                data_preview = []
                for row in rows:
                    data_preview.append(dict(zip(columns, row)))
                
                # Generate simple answer based on query content
                if "sales" in original_query.lower() or "trend" in original_query.lower():
                    total_amount = sum(float(row.get('total_amount', 0)) for row in data_preview if row.get('total_amount'))
                    return f"Based on recent data from {table_name}, I found {len(data_preview)} records with a total amount of approximately {total_amount:.2f}. Due to processing limitations, this is a simplified analysis."
                elif "count" in original_query.lower() or "how many" in original_query.lower():
                    count_query = f"SELECT COUNT(*) as count FROM {table_name}"
                    count_result = connection.execute(text(count_query))
                    count = count_result.fetchone()[0]
                    return f"The table {table_name} contains {count} total records."
                else:
                    return f"I found {len(data_preview)} sample records in {table_name}. Here's a preview of the data structure."
            else:
                return f"The table {table_name} appears to be empty or inaccessible."
                
    except Exception as e:
        logger.error(f"Fallback query also failed: {e}")
        return "I encountered processing difficulties. Please try simplifying your question or contact support."

def initialize_app_state():
    """
    Initializes critical application state, like database and LLM/Embedding models.
    Called at application startup.
    """
    logger.info("Starting initialization of application state (database, LLM, embedding model)...")
    
    # 1. Initialize Database (ensure schema and base data)
    try:
        initialize_database() # Changed: Call the renamed function in db.py
        logger.info("Database initialization completed.")
    except Exception as e:
        logger.error(f"An error occurred during database initialization: {e}", exc_info=True)
        # Depending on severity, might want to raise to stop app, or continue with limited functionality.

    # LLM and Embeddings are already initialized when this module is imported.
    # We can add checks here or re-log their status.
    if llm:
        logger.info("LLM model was initialized when this module was loaded.")
    else:
        logger.warning("LLM model not initialized. Question answering and report functionality will be limited.")

    if embeddings:
        logger.info("Embedding model was initialized when this module was loaded.")
    else:
        logger.warning("Embedding model not initialized. RAG functionality will be unavailable.")
        
    logger.info("Application state initialization completed.")

# Add new function for query SQL processing
async def get_query_from_sqltable_datasource(
    query: str, 
    active_datasource: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Get answer from SQL table datasource with improved error handling and SQL cleaning
    Specifically optimized for data queries (not charts)
    """
    logger.info(f"Attempting Query SQL on datasource: {active_datasource['name']} (ID: {active_datasource['id']}) for query: '{query}'")

    if not llm:
        return {
            "success": False,
            "error": "LLM not initialized. SQL query cannot be processed.",
            "query": query
        }

    try:
        # Get table name from datasource
        table_name = f"dstable_{active_datasource['id']}_sample_sales_2025_19f7aeb7"
        logger.info(f"Initializing SQLDatabase for query: {table_name}")
        
        # Initialize database connection
        db = SQLDatabase.from_uri(DB_URI)
        
        # Get schema info for validation
        schema_info = db.get_table_info()
        
        # Process time expressions in query
        processed_query = _process_time_expressions(query)

        # Get available tables
        available_tables = db.get_usable_table_names()
        logger.info(f"Available tables: {available_tables}")

        # Generate SQL using direct LLM call
        sql_generation_prompt = f"""
        User query: "{processed_query}"

        Available tables: {available_tables}
        Table schema: {schema_info}

        Please generate a SQLite query to answer the user's question. The query should:
        1. Only use SELECT statements (no INSERT, UPDATE, DELETE, etc.)
        2. Only use tables that are available in the database
        3. Use appropriate time filters for temporal queries:
           - "this month" -> WHERE strftime('%Y-%m', saledate) = strftime('%Y-%m', 'now')
           - "last month" -> WHERE strftime('%Y-%m', saledate) = strftime('%Y-%m', 'now', '-1 month')
           - "this year" -> WHERE strftime('%Y', saledate) = strftime('%Y', 'now')
        4. Use appropriate aggregations (SUM, COUNT, AVG) if needed
        5. Include proper WHERE clauses to filter data
        6. Use GROUP BY if needed to aggregate data
        7. Order results appropriately
        8. Limit results to a reasonable number (e.g., LIMIT 50)

        Only return the SQL query, no explanations or other text.
        """

        try:
            # Execute query with timeout from config
            sql_response = await asyncio.wait_for(
                llm.agenerate([sql_generation_prompt]),
                timeout=Config.LLM_TIMEOUT
            )
            
            # Extract SQL from response
            sql = sql_response.generations[0][0].text.strip()
            
            # Clean and validate SQL
            clean_sql = _clean_sql_statement(sql)
            if not _validate_sql_statement(clean_sql, schema_info):
                raise ValueError(f"Invalid SQL statement generated: {clean_sql}")
            
            try:
                # Execute the cleaned SQL
                query_result = db.run(clean_sql)
                
                # Handle different result types
                if isinstance(query_result, str):
                    # If result is a string, it might be an error or a single value
                    logger.warning(f"Query returned string result: {query_result}")
                    
                    # Try to parse if it's a string representation of data
                    parsed_data = _parse_string_query_result(query_result)
                    if parsed_data:
                        structured_data = {
                            "rows": parsed_data,
                            "columns": ["label", "value"] if len(parsed_data[0]) == 2 else [f"col_{i}" for i in range(len(parsed_data[0]))],
                            "executed_sql": clean_sql,
                            "queried_table": table_name
                        }
                    else:
                        # Fallback to original logic for simple string results
                        structured_data = {
                            "rows": [{"result": query_result}],
                            "columns": ["result"],
                            "executed_sql": clean_sql,
                            "queried_table": table_name
                        }
                else:
                    # Normal case: result is a list of dictionaries
                    if not query_result:
                        return {
                            "success": False,
                            "error": "Query returned no results. Please check if the time period contains data.",
                            "query": processed_query,
                            "executed_sql": clean_sql
                        }
                        
                    # Ensure query_result is a list
                    if not isinstance(query_result, list):
                        query_result = [query_result]
                        
                    structured_data = {
                        "rows": query_result,
                        "columns": list(query_result[0].keys()) if query_result else [],
                        "executed_sql": clean_sql,
                        "queried_table": table_name
                    }
                
                return {
                    "success": True,
                    "data": structured_data,
                    "answer": f"Here are the results for your query about {processed_query}",
                    "executed_sql": clean_sql,
                    "query": processed_query
                }
                
            except Exception as e:
                error_msg = str(e)
                if "no such table" in error_msg.lower():
                    error_msg = f"Table not found. Available tables are: {', '.join(available_tables)}"
                elif "no such column" in error_msg.lower():
                    error_msg = f"Column not found in schema: {error_msg}"
                    
                logger.error(f"SQL execution failed: {error_msg}")
                return {
                    "success": False,
                    "error": f"Error executing SQL query: {error_msg}",
                    "query": processed_query,
                    "executed_sql": clean_sql
                }

        except asyncio.TimeoutError:
            logger.warning(f"Query SQL generation timed out for query: {processed_query}")
            return {
                "success": False,
                "error": "Query generation timed out",
                "query": processed_query
            }

    except Exception as e:
        logger.error(f"Error in get_query_from_sqltable_datasource: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Failed to process query: {str(e)}",
            "query": query
        }

def _clean_sql_statement(sql: str) -> str:
    """Clean SQL statement by removing comments and extra text"""
    # Remove comments
    sql = re.sub(r'/\*.*?\*/', '', sql, flags=re.DOTALL)  # Remove /* ... */ comments
    sql = re.sub(r'--.*$', '', sql, flags=re.MULTILINE)   # Remove -- comments
    
    # Remove explanation text after semicolon
    sql = sql.split(';')[0] + ';' if ';' in sql else sql
    
    # Clean whitespace
    sql = ' '.join(sql.split())
    
    return sql

def _validate_sql_statement(sql: str, schema_info: str) -> bool:
    """Validate SQL statement against schema and basic rules"""
    sql_lower = sql.lower()
    
    # Basic security checks
    dangerous_keywords = ['drop', 'truncate', 'delete', 'update', 'insert', 'alter', 'create']
    if any(keyword in sql_lower for keyword in dangerous_keywords):
        return False
    
    # Must be a SELECT statement
    if not sql_lower.strip().startswith('select'):
        return False
    
    return True

def _extract_sql_from_result(result: Dict[str, Any]) -> Optional[str]:
    """Extract SQL statement from agent result"""
    if not result or not isinstance(result, dict):
        return None
        
    # Try to find SQL in the output
    output = result.get("output", "")
    if not output:
        return None
        
    # Look for SQL between backticks or after "SQL:" or similar markers
    sql_markers = [
        r"```sql\n(.*?)\n```",
        r"```(.*?)```",
        r"SQL:\s*(.*?)(?:\n|$)",
        r"Query:\s*(.*?)(?:\n|$)",
        r"Generated SQL:\s*(.*?)(?:\n|$)"
    ]
    
    for marker in sql_markers:
        match = re.search(marker, output, re.DOTALL | re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    # If no SQL found in markers, try to find the first statement that looks like SQL
    lines = output.split('\n')
    for line in lines:
        if line.strip().lower().startswith('select'):
            return line.strip()
    
    return None

def _process_time_expressions(query: str) -> str:
    """Process time-related expressions in the query"""
    # Convert "last X weeks/months/years" to specific date ranges
    patterns = [
        (r'last (\d+) weeks?', lambda m: f'in the past {m.group(1)} weeks'),
        (r'last (\d+) months?', lambda m: f'in the past {m.group(1)} months'),
        (r'last (\d+) years?', lambda m: f'in the past {m.group(1)} years'),
        (r'past (\d+) weeks?', lambda m: f'in the past {m.group(1)} weeks'),
        (r'past (\d+) months?', lambda m: f'in the past {m.group(1)} months'),
        (r'past (\d+) years?', lambda m: f'in the past {m.group(1)} years'),
        (r'previous (\d+) weeks?', lambda m: f'in the past {m.group(1)} weeks'),
        (r'previous (\d+) months?', lambda m: f'in the past {m.group(1)} months'),
        (r'previous (\d+) years?', lambda m: f'in the past {m.group(1)} years')
    ]
    
    processed = query
    for pattern, replacement in patterns:
        processed = re.sub(pattern, replacement, processed, flags=re.IGNORECASE)
    
    return processed

def _parse_string_query_result(result: str) -> Optional[List[Dict[str, Any]]]:
    """Parse a string result that contains Python tuple list representation"""
    try:
        # Check if the result looks like a Python tuple list string
        if result.startswith('[') and result.endswith(']') and '(' in result and ')' in result:
            # Try to safely evaluate the string as Python literal
            import ast
            try:
                # Use ast.literal_eval for safe parsing of Python literals
                parsed_tuples = ast.literal_eval(result)
                
                if isinstance(parsed_tuples, list) and all(isinstance(item, tuple) for item in parsed_tuples):
                    # Convert tuples to dictionaries with generic column names
                    parsed_data = []
                    for tuple_item in parsed_tuples:
                        if len(tuple_item) == 2:
                            # Common case: (label, value) pairs
                            parsed_data.append({
                                "label": str(tuple_item[0]),
                                "value": float(tuple_item[1]) if isinstance(tuple_item[1], (int, float)) else tuple_item[1]
                            })
                        else:
                            # Multiple columns: create numbered columns
                            row_dict = {}
                            for i, value in enumerate(tuple_item):
                                row_dict[f"col_{i}"] = value
                            parsed_data.append(row_dict)
                    
                    logger.info(f"Successfully parsed {len(parsed_data)} rows from string result")
                    return parsed_data
                    
            except (ValueError, SyntaxError) as e:
                logger.warning(f"Failed to parse tuple list string: {e}")
                return None
        
        # If not a tuple list format, return None to use fallback logic
        return None
        
    except Exception as e:
        logger.error(f"Error parsing string query result: {e}")
        return None

